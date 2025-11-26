import streamlit as st
import pandas as pd
import numpy as np
import re
import io
import zipfile
import sqlite3
from datetime import datetime
from typing import Optional, List, Tuple
import logging

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ==========================
# Configuração de Logging
# ==========================
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


# ==========================
# Classe para gerenciar histórico
# ==========================
class HistoricoRepository:
    def __init__(self, db_path: str = "historico_relatorios.db"):
        self.db_path = db_path
        self.conn = sqlite3.connect(db_path, check_same_thread=False)
        self._criar_tabela()
    
    def _criar_tabela(self):
        """Cria a tabela de histórico se não existir."""
        try:
            cursor = self.conn.cursor()
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS historico (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    data_hora TEXT,
                    nome_trilha TEXT,
                    nome_arquivo_xlsx TEXT,
                    abas_selecionadas TEXT,
                    qtd_relatorios INTEGER
                )
            """)
            self.conn.commit()
            logger.info("Tabela de histórico verificada/criada com sucesso")
        except Exception as e:
            logger.error(f"Erro ao criar tabela: {e}")
            raise
    
    def registrar(self, nome_trilha: str, nome_arquivo_xlsx: str, 
                  abas: List[str], qtd_relatorios: int):
        """Registra um novo histórico de geração."""
        try:
            data_hora = datetime.now().isoformat(timespec="seconds")
            abas_str = ", ".join(abas)
            cursor = self.conn.cursor()
            cursor.execute(
                """
                INSERT INTO historico (data_hora, nome_trilha, nome_arquivo_xlsx, 
                                     abas_selecionadas, qtd_relatorios)
                VALUES (?, ?, ?, ?, ?)
                """,
                (data_hora, nome_trilha, nome_arquivo_xlsx, abas_str, qtd_relatorios)
            )
            self.conn.commit()
            logger.info(f"Histórico registrado: {nome_trilha} - {qtd_relatorios} relatórios")
        except Exception as e:
            logger.error(f"Erro ao registrar histórico: {e}")
            st.error(f"Erro ao salvar histórico: {str(e)}")
    
    def obter_historico(self) -> pd.DataFrame:
        """Retorna o histórico completo."""
        try:
            return pd.read_sql_query(
                "SELECT * FROM historico ORDER BY id DESC", 
                self.conn
            )
        except Exception as e:
            logger.error(f"Erro ao obter histórico: {e}")
            return pd.DataFrame()
    
    def limpar_historico(self) -> bool:
        """Limpa todo o histórico. Retorna True se bem-sucedido."""
        try:
            cursor = self.conn.cursor()
            cursor.execute("DELETE FROM historico")
            self.conn.commit()
            logger.info("Histórico limpo com sucesso")
            return True
        except Exception as e:
            logger.error(f"Erro ao limpar histórico: {e}")
            return False
    
    def fechar(self):
        """Fecha a conexão com o banco."""
        if self.conn:
            self.conn.close()
            logger.info("Conexão com banco fechada")


# ==========================
# Funções de validação
# ==========================
def validar_nome_trilha(nome: str) -> Tuple[bool, str]:
    """
    Valida o nome da trilha.
    Retorna (válido, mensagem_erro)
    """
    if not nome or len(nome.strip()) < 3:
        return False, "O nome da trilha deve ter pelo menos 3 caracteres."
    
    # Caracteres inválidos para nomes de arquivo
    if re.search(r'[<>:"/\\|?*]', nome):
        return False, "O nome contém caracteres inválidos (< > : \" / \\ | ? *)."
    
    return True, ""


def validar_estrutura_excel(df: pd.DataFrame) -> Tuple[bool, str]:
    """
    Valida se o DataFrame possui a estrutura esperada.
    Retorna (válido, mensagem_erro)
    """
    if df.empty:
        return False, "A aba está vazia."
    
    if len(df) < 2:
        return False, "A aba não possui dados suficientes (mínimo 2 linhas)."
    
    return True, ""


# ==========================
# Funções auxiliares de negócio
# ==========================
def aba_e_numerica(nome: str) -> bool:
    """Verifica se o nome da aba começa com número."""
    return re.match(r"^\d", str(nome).strip()) is not None


def formatar_celula(celula, bold=False, branca=False):
    """Aplica formatação Arial 12 na célula."""
    for par in celula.paragraphs:
        for run in par.runs:
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.bold = bold
            if branca:
                run.font.color.rgb = RGBColor(255, 255, 255)


def pintar_fundo(celula, rgb=(0, 0, 0)):
    """Aplica cor de fundo (w:shd) na célula."""
    cor = "%02X%02X%02X" % rgb

    tc = celula._tc
    tcPr = tc.get_or_add_tcPr()

    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)

    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), cor)


def gerar_doc_para_aba(xls: pd.ExcelFile, aba_nome: str) -> Optional[bytes]:
    """
    Gera o DOCX (em bytes) para uma aba específica.
    Retorna None se houver erro ou dados insuficientes.
    """
    try:
        df = pd.read_excel(xls, sheet_name=aba_nome, header=None)
        
        # Validar estrutura
        valido, msg_erro = validar_estrutura_excel(df)
        if not valido:
            logger.warning(f"Aba {aba_nome}: {msg_erro}")
            return None

        header = df.iloc[1]
        
        # Buscar colunas "Termo" e "Campo"
        try:
            col_termo = header[header.astype(str).str.contains("Termo", case=False)].index[0]
            col_campo = header[header.astype(str).str.contains("Campo", case=False)].index[0]
        except IndexError:
            logger.warning(f"Aba {aba_nome}: Colunas 'Termo' ou 'Campo' não encontradas")
            return None

        dados = df.iloc[2:, [col_termo, col_campo]].copy()
        dados.columns = ["Termo", "Campo"]
        dados = dados.replace({np.nan: None})
        dados = dados[~(dados["Termo"].isnull() & dados["Campo"].isnull())]

        if dados.empty:
            logger.warning(f"Aba {aba_nome}: Sem dados para processar")
            return None

        # Criar documento em paisagem
        doc = Document()
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        doc.add_paragraph(f"Relatório – {aba_nome}")

        # Criar tabela
        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Cabeçalho
        hdr = table.rows[0].cells
        hdr[0].text = "Termo:"
        hdr[1].text = "Campo de preenchimento:"
        hdr[2].text = "Evidência"

        for c in hdr:
            pintar_fundo(c, rgb=(0, 0, 0))
            formatar_celula(c, bold=True, branca=True)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adicionar linhas de dados
        for _, row in dados.iterrows():
            termo = row["Termo"] if row["Termo"] else ""
            campo = row["Campo"] if row["Campo"] else ""

            nova = table.add_row().cells
            nova[0].text = str(termo)
            nova[1].text = str(campo)
            nova[2].text = ""

            for c in nova:
                formatar_celula(c)

        # Salvar em buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"Documento gerado com sucesso para aba: {aba_nome}")
        return buffer.getvalue()
        
    except Exception as e:
        logger.error(f"Erro ao gerar documento para aba {aba_nome}: {e}")
        st.error(f"Erro ao processar aba '{aba_nome}': {str(e)}")
        return None


# ==========================
# Configuração da página
# ==========================
def configurar_pagina():
    """Configura a página e aplica estilos CSS."""
    st.set_page_config(page_title="Gerador de Relatórios TIC", layout="wide")
    st.title("Gerador de Relatórios TIC → Tabelas em Word")
    
    st.markdown(
    """
    <style>
    
    .stApp {
        background-color: #E6F0FA !important;
    }

    h1, h2, h3, h4 {
        color: #003F88 !important;
        font-weight: 700 !important;
    }

    p, label, span {
        color: #1C1C1C !important;
    }

    .stFileUploader label {
        color: #003F88 !important;
        font-weight: 600 !important;
    }

    /* ------------------------------
       🔵 Botões Azul (normal)
       ------------------------------ */
    .stButton > button {
        background-color: #003F88 !important;
        color: white !important;
        border-radius: 8px !important;
        padding: 0.6rem 1.2rem !important;
        font-weight: 600 !important;
        border: 2px solid #003F88 !important;
        transition: all 0.3s ease !important;
    }

    /* ------------------------------
       🟠 Botões em hover (laranja)
       ------------------------------ */
    .stButton > button:hover {
        background-color: #F28C28 !important;
        border-color: #F28C28 !important;
        color: white !important; /* <=== FORÇANDO TEXTO BRANCO NO HOVER */
    }

    /* ------------------------------
       🟠 Botões ativos/focados
       ------------------------------ */
    .stButton > button:active,
    .stButton > button:focus {
        background-color: #F28C28 !important;
        border-color: #F28C28 !important;
        color: white !important; /* <=== TAMBÉM BRANCO AQUI */
    }

    /* Caixas de texto */
    input[type="text"] {
        border: 2px solid #003F88 !important;
        border-radius: 6px !important;
    }

    /* Multiselect */
    .stMultiSelect > div {
        border: 2px solid #003F88 !important;
        border-radius: 6px !important;
    }

    </style>
    """,
    unsafe_allow_html=True
)



@st.cache_resource
def carregar_excel(_uploaded_file) -> pd.ExcelFile:
    """Carrega o arquivo Excel com cache."""
    return pd.ExcelFile(_uploaded_file)


# ==========================
# Aplicação principal
# ==========================
def main():
    configurar_pagina()
    
    # Inicializar repositório de histórico
    historico = HistoricoRepository()
    
    # Upload do arquivo
    uploaded_file = st.file_uploader(
        "Envie o arquivo de parametrização TIC (.xlsx)",
        type=["xlsx"],
        help="Arquivo Excel contendo as abas numeradas com os dados dos relatórios"
    )

    if uploaded_file is not None:
        # Input do nome da trilha
        nome_trilha = st.text_input(
            "Nome da trilha (usado no nome do arquivo ZIP):",
            placeholder="Ex.: Trilha_IA, TIC_Módulo_3, Material_Senac",
            help="Mínimo 3 caracteres, sem caracteres especiais"
        )

        try:
            # Carregar Excel com cache
            xls = carregar_excel(uploaded_file)
            abas_numericas = [a for a in xls.sheet_names if aba_e_numerica(a)]

            if not abas_numericas:
                st.warning("⚠️ Nenhuma aba numerada encontrada no arquivo.")
            else:
                st.markdown("### Abas numeradas detectadas")
                st.info(f"📊 {len(abas_numericas)} aba(s) encontrada(s)")
                
                # Seleção de abas
                aba_sel = st.multiselect(
                    "Selecione as abas para gerar relatório:",
                    options=abas_numericas,
                    default=abas_numericas,
                    help="Selecione uma ou mais abas para processar"
                )

                # Botão de gerar
                if st.button("🚀 Gerar relatórios (DOCX)"):
                    # Validar nome da trilha
                    valido, msg_erro = validar_nome_trilha(nome_trilha)
                    if not valido:
                        st.error(f"⚠️ {msg_erro}")
                        st.stop()

                    if not aba_sel:
                        st.error("⚠️ Selecione pelo menos uma aba.")
                        st.stop()

                    # Nome final do ZIP
                    nome_zip_base = nome_trilha.strip()
                    nome_zip = nome_zip_base + ".zip"

                    # Barra de progresso
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    zip_buffer = io.BytesIO()
                    qtd_ok = 0
                    erros = []

                    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
                        for idx, aba in enumerate(aba_sel):
                            status_text.text(f"Processando: {aba}...")
                            
                            doc_bytes = gerar_doc_para_aba(xls, aba)
                            
                            if doc_bytes is None:
                                erros.append(aba)
                            else:
                                nome_arquivo = f"Tabela_{aba}.docx"
                                zipf.writestr(nome_arquivo, doc_bytes)
                                qtd_ok += 1
                            
                            # Atualizar progresso
                            progress_bar.progress((idx + 1) / len(aba_sel))

                    progress_bar.empty()
                    status_text.empty()
                    
                    zip_buffer.seek(0)

                    # Registrar no histórico
                    historico.registrar(
                        nome_trilha=nome_trilha.strip(),
                        nome_arquivo_xlsx=uploaded_file.name,
                        abas=aba_sel,
                        qtd_relatorios=qtd_ok
                    )

                    # Exibir resultado
                    if qtd_ok > 0:
                        st.success(f"✅ Relatórios gerados com sucesso! ({qtd_ok} arquivo(s))")
                        
                        if erros:
                            st.warning(f"⚠️ {len(erros)} aba(s) não puderam ser processadas: {', '.join(erros)}")
                        
                        st.download_button(
                            label="📥 Baixar ZIP com relatórios",
                            data=zip_buffer,
                            file_name=nome_zip,
                            mime="application/zip"
                        )
                    else:
                        st.error("❌ Nenhum relatório pôde ser gerado. Verifique a estrutura das abas.")
                        if erros:
                            st.info(f"Abas com erro: {', '.join(erros)}")

        except Exception as e:
            logger.error(f"Erro ao processar arquivo: {e}")
            st.error(f"❌ Erro ao processar o arquivo: {str(e)}")

    else:
        st.info("📤 Envie o arquivo Excel para começar.")

    # ==========================
    # Seção de histórico
    # ==========================
    st.markdown("---")
    st.subheader("📜 Histórico de relatórios gerados")

    if st.checkbox("Mostrar histórico", help="Exibe todos os relatórios já gerados"):
        df_hist = historico.obter_historico()
        if df_hist.empty:
            st.info("Ainda não há registros no histórico.")
        else:
            # Criar colunas para botões
            col1, col2 = st.columns([1, 5])
            
            with col1:
                if st.button("🗑️ Limpar histórico", type="secondary"):
                    if historico.limpar_historico():
                        st.success("✅ Histórico limpo com sucesso!")
                        st.rerun()
                    else:
                        st.error("❌ Erro ao limpar histórico")
            
            with col2:
                # Botão para exportar histórico
                csv = df_hist.to_csv(index=False).encode('utf-8')
                st.download_button(
                    label="📊 Exportar histórico como CSV",
                    data=csv,
                    file_name=f"historico_relatorios_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                    mime="text/csv"
                )
            
            st.dataframe(df_hist, use_container_width=True)


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        logger.error(f"Erro fatal na aplicação: {e}")
        st.error(f"Ocorreu um erro inesperado: {str(e)}")